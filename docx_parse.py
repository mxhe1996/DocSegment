import docx.text.paragraph
from docx import Document
import datetime
import copy
from semantic_seg import SemanticTextSplitter
from text2vec import SimpleEmbeddingModel

doc_path = f"./demo/test1.docx"
doc = Document(doc_path)


spliter = SemanticTextSplitter(
    embed_model=SimpleEmbeddingModel())


def build_struct_tree(charter_content: list[docx.text.paragraph.Paragraph], character_lv: int, _meta: dict):
    # assert charter_content[0].style.name == 'Heading ' + str(character_lv)
    nodes = []
    level_index = []
    for i, single_para in enumerate(charter_content):
        if single_para.style.name.startswith('Heading '+str(character_lv)):
            level_index.append(i)

    if len(level_index) == 0:
        semantic_chunk = spliter.semantic_split_text(''.join([p.text for p in charter_content]).strip())
        # semantic_chunk = ''.join([p.text for p in charter_content])
        return {'content': semantic_chunk, '_meta': _meta}, ''  # 同一文本内容分段

    for i in range(len(level_index)):
        body = {'title': charter_content[level_index[i]].text}
        inner_meta = copy.deepcopy(_meta)
        inner_meta['hierarchy'] = inner_meta['hierarchy']+"/"+body['title']
        start_flag = level_index[i]+1
        end_flag = level_index[i+1] if i+1 < len(level_index) else -1
        content, children_array = build_struct_tree(charter_content[start_flag:end_flag], character_lv+1, inner_meta)
        body['children'] = children_array
        body['content'] = content
        nodes.append(body)

    pre_content = ''.join([p.text for p in charter_content[:level_index[0]]]).strip()
    return {'content': spliter.semantic_split_text(pre_content), '_meta': _meta}, nodes
    # return {'content': pre_content, '_meta': _meta}, nodes


for index, para in enumerate(doc.paragraphs):
    style_ = para.style.name
    if style_.startswith('Heading '):
        head_lv = int(style_.replace('Heading ', ''))
        print('\t'*head_lv+para.text)

meta = {
    'doc_name': doc_path.split("\\")[-1],
    'build_time': datetime.datetime.now().strftime("%Y-%M-%d %H:%M:%S"),
    'hierarchy': doc_path.split("\\")[-1][:doc_path.split("\\")[-1].index('.')]
}
_, tree = build_struct_tree(doc.paragraphs, 1, meta)


import json
tree_json = json.dumps(tree)
with open('./out.txt', 'w+') as out_file:
    out_file.write(tree_json)

